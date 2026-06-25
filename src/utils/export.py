"""
Utilities to render Markdown content into PDF or DOCX,
without relying on external Markdown or WeasyPrint libraries.
"""

import os

import tempfile

import sys

from html import escape

import re

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document

from reportlab.lib.pagesizes import LETTER

from reportlab.lib.styles import getSampleStyleSheet

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


def render_markdown_to_docx(markdown_text: str) -> bytes:
    """
    Convert Markdown text to DOCX bytes.
    Basic implementation: each line becomes a paragraph.
    """

    doc = Document()

    for line in markdown_text.splitlines():

        stripped = line.strip()

        if not stripped:

            doc.add_paragraph()

        elif stripped.startswith("#"):

            level = min(9, len(stripped) - len(stripped.lstrip("#")))

            doc.add_heading(stripped.lstrip("# "), level=level)

        elif re.match(r"^[-*]\s+", stripped):

            doc.add_paragraph(re.sub(r"^[-*]\s+", "", stripped), style="List Bullet")

        else:

            doc.add_paragraph(stripped)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:

        tmp_path = tmp.name

    try:

        doc.save(tmp_path)

        with open(tmp_path, "rb") as f:

            content = f.read()

    finally:

        os.remove(tmp_path)

    return content


def render_markdown_to_pdf(markdown_text: str) -> bytes:
    """
    Convert Markdown text to PDF bytes, using ReportLab.
    Basic implementation: each line becomes a paragraph.
    """

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:

        tmp_path = tmp.name

    try:

        doc = SimpleDocTemplate(tmp_path, pagesize=LETTER)

        styles = getSampleStyleSheet()

        flowables = []

        for line in markdown_text.splitlines():

            stripped = line.strip()

            if not stripped:

                flowables.append(Spacer(1, 12))

            else:

                heading_level = len(stripped) - len(stripped.lstrip("#"))

                text = stripped.lstrip("# ") if heading_level else stripped

                text = re.sub(r"^[-*]\s+", "• ", text)

                style = (
                    styles.get(f"Heading{min(heading_level, 3)}", styles["Normal"])
                    if heading_level
                    else styles["Normal"]
                )

                para = Paragraph(escape(text), style)

                flowables.append(para)

        doc.build(flowables)

        with open(tmp_path, "rb") as f:

            pdf_bytes = f.read()

    finally:

        os.remove(tmp_path)

    return pdf_bytes
